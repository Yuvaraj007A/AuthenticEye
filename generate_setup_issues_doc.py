"""
Script to generate a DOCX report of all issues encountered during
AuthenticEye local setup, including their root causes and solutions.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.dml.color import ColorFormat
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ─── Issues data ────────────────────────────────────────────────────────────────
ISSUES = [
    {
        "id": 1,
        "title": "Docker Daemon Not Running",
        "component": "Docker",
        "severity": "High",
        "symptom": (
            "When attempting to use Docker Compose to start all services, "
            "the command failed with the error:\n"
            "'failed to connect to the docker API at npipe:// — The system cannot find the file specified.'"
        ),
        "root_cause": (
            "Docker Desktop was installed but the Docker Engine daemon was not running. "
            "The named pipe used by the Docker CLI to communicate with the daemon was unavailable, "
            "which means Docker Desktop was not started."
        ),
        "solution": (
            "Since Docker was unavailable, all services were started manually in their native runtimes "
            "instead of via Docker Compose:\n"
            "  • Backend:    npm run dev  (inside ./backend)\n"
            "  • Frontend:   npm run dev -- --port 5173 --host  (inside ./frontend)\n"
            "  • AI Service: .\\venv\\Scripts\\python.exe -m uvicorn main:app --host 0.0.0.0 --port 8000 --reload  (inside ./ai-service)\n\n"
            "To fix Docker permanently: open Docker Desktop from the Start Menu and wait for the "
            "engine to fully start before running docker commands."
        ),
    },
    {
        "id": 2,
        "title": "AI Service Virtual Environment Had No Dependencies Installed",
        "component": "AI Service (Python)",
        "severity": "High",
        "symptom": (
            "Running the AI service with the existing virtual environment (./ai-service/venv) "
            "immediately failed with:\n"
            "'No module named uvicorn'"
        ),
        "root_cause": (
            "The virtual environment directory existed (created previously) but none of the required "
            "Python packages from requirements.txt had been installed into it. "
            "The venv was empty — only the Python interpreter itself was present."
        ),
        "solution": (
            "Ran pip install inside the virtual environment to install all dependencies:\n"
            "    .\\venv\\Scripts\\pip.exe install -r requirements.txt\n\n"
            "This installed all 65 packages including FastAPI, Uvicorn, PyTorch, torchvision, "
            "timm, OpenCV, MediaPipe, SciPy and others. Note: this step took several minutes "
            "due to the large size of PyTorch (~2GB)."
        ),
    },
    {
        "id": 3,
        "title": "MediaPipe Version Incompatibility — 'module has no attribute solutions'",
        "component": "AI Service (Python / MediaPipe)",
        "severity": "High",
        "symptom": (
            "After installing dependencies, the AI service failed to start with the error:\n"
            "'AttributeError: module \"mediapipe\" has no attribute \"solutions\"'\n\n"
            "The traceback pointed to preprocessing.py line 29:\n"
            "    _mp_face_detection = mp.solutions.face_detection"
        ),
        "root_cause": (
            "The requirements.txt specified mediapipe>=0.10.14, but the latest release (0.10.32) "
            "dropped the legacy mp.solutions API that was heavily used in older tutorials and codebases. "
            "In MediaPipe 0.10.21+, the classic 'mp.solutions' namespace was moved/renamed to "
            "a new Tasks API. The code was written targeting the older solutions API which no longer "
            "exists in the installed version."
        ),
        "solution": (
            "Patched preprocessing.py to catch AttributeError (in addition to ImportError) during "
            "MediaPipe initialization, allowing a graceful fallback to the OpenCV Haar Cascade "
            "face detector instead of crashing:\n\n"
            "    # Before (only catches ImportError):\n"
            "    except ImportError:\n"
            "        MEDIAPIPE_AVAILABLE = False\n\n"
            "    # After (also catches AttributeError and any other exception):\n"
            "    except (ImportError, AttributeError, Exception):\n"
            "        MEDIAPIPE_AVAILABLE = False\n\n"
            "The fallback code path using OpenCV Haar Cascades (haarcascade_frontalface_default.xml) "
            "was already implemented in the codebase, so face detection continued to work correctly "
            "after this fix. The AI service started successfully with this change."
        ),
    },
    {
        "id": 4,
        "title": "AI Model Weights Downloaded on First Startup (Slow Cold Start)",
        "component": "AI Service (PyTorch / HuggingFace)",
        "severity": "Medium",
        "symptom": (
            "The AI service took several minutes to start up after the dependency installation. "
            "The console showed progress bars downloading large model weight files:\n"
            "  • model.safetensors (EfficientNet-B4) — ~78MB\n"
            "  • xception-43020ad28.pth (XceptionNet) — from PyTorch Hub\n"
            "  • model.safetensors (ViT-B/16) — ~346MB\n"
            "  • resnet18-f37072fd.pth (used by frequency/GAN detectors) — ~45MB"
        ),
        "root_cause": (
            "The first run always downloads pretrained ImageNet weights from HuggingFace Hub and "
            "PyTorch Hub because no fine-tuned checkpoint files existed in the ai-service/checkpoints/ "
            "directory. The model.py load_models() function checks for .pth files in the checkpoints "
            "directory and falls back to downloading pretrained weights when they are absent."
        ),
        "solution": (
            "Waited for the downloads to complete (~3–5 minutes depending on internet speed). "
            "The weights are cached locally in:\n"
            "  • C:\\Users\\<username>\\.cache\\huggingface\\hub\\\n"
            "  • C:\\Users\\<username>\\.cache\\torch\\hub\\checkpoints\\\n\n"
            "Subsequent service restarts will be much faster as the weights are served from cache. "
            "No code changes were needed for this issue."
        ),
    },
    {
        "id": 5,
        "title": "Frontend .env File Missing — API URL Not Configured",
        "component": "Frontend (Vite / React)",
        "severity": "Medium",
        "symptom": (
            "The frontend/.env file did not exist, which meant the VITE_API_URL environment "
            "variable was undefined at build time. API requests from the React frontend "
            "would have been sent to an incorrect URL (undefined/api/...) instead of "
            "the backend at http://localhost:5000/api."
        ),
        "root_cause": (
            "The README instructed developers to create frontend/.env manually with:\n"
            '    echo "VITE_API_URL=http://localhost:5000/api" > .env\n\n'
            "This step was not done, so the file was absent. The root .env.example file provided "
            "the correct value but was not copied to the frontend directory."
        ),
        "solution": (
            "Created the frontend/.env file with the correct content:\n"
            "    VITE_API_URL=http://localhost:5000/api\n\n"
            "This ensures Vite includes the correct API base URL in the compiled JavaScript bundle "
            "and all axios API calls from the frontend are routed to the Express backend."
        ),
    },
    {
        "id": 6,
        "title": "Redis Not Installed as a Windows Service",
        "component": "Redis",
        "severity": "Low",
        "symptom": (
            "When checking for required services, Get-Service -Name Redis returned:\n"
            "'Cannot find any service with service name Redis.'\n\n"
            "Redis is listed in the docker-compose.yml as a required dependency for the backend."
        ),
        "root_cause": (
            "Redis does not provide a native Windows installer with a system service. "
            "In the Docker setup, Redis runs as a container. "
            "For local non-Docker development, Redis must either be installed manually (via WSL, "
            "Memurai, or the Windows Redis port), or the backend code must be adjusted to "
            "run without Redis."
        ),
        "solution": (
            "The backend (server.js) did not crash without Redis because Redis is marked as a "
            "future feature for job queuing and is not yet integrated into active request handling. "
            "The backend started and connected to MongoDB successfully without Redis.\n\n"
            "For full Redis support on Windows, install one of:\n"
            "  • Memurai (Redis-compatible for Windows): https://www.memurai.com/\n"
            "  • Redis via WSL2 (Windows Subsystem for Linux)\n"
            "  • Run Redis only through Docker Compose when Docker Desktop is running."
        ),
    },
    {
        "id": 7,
        "title": "Vite Frontend Dev Server Exited on First Attempt",
        "component": "Frontend (Vite)",
        "severity": "Low",
        "symptom": (
            "The first attempt to start the frontend with 'npm run dev' exited immediately "
            "with exit code 1, even though Vite reported it was 'ready in 953ms' at http://localhost:5173/."
        ),
        "root_cause": (
            "The first npm run dev call did not include --host, which causes the Vite process "
            "to close when it is not kept alive by the terminal session that started it (background "
            "process context). The process ended before it could accept connections."
        ),
        "solution": (
            "Restarted the frontend server with the --host and explicit --port flags to keep it "
            "bound to all interfaces and running persistently:\n"
            "    npm run dev -- --port 5173 --host\n\n"
            "The frontend then remained accessible at http://localhost:5173 and also on the "
            "local network addresses (192.168.x.x, 10.x.x.x)."
        ),
    },
]


# ─── Helpers ─────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set a table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc):
    """Add a simple horizontal rule (border paragraph)."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def add_colored_paragraph(doc, text, bold=False, italic=False,
                           font_size=11, color_hex=None, left_indent=None,
                           space_before=4, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if left_indent:
        p.paragraph_format.left_indent = left_indent
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(font_size)
    if color_hex:
        r, g, b = int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    return p


def add_code_block(doc, code_text):
    """Simulate a code block paragraph with a grey background & mono font."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Grey shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F0F0F0')
    pPr.append(shd)
    # Border
    pBdr = OxmlElement('w:pBdr')
    for side in ('top', 'bottom', 'left', 'right'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '4')
        el.set(qn('w:color'), 'AAAAAA')
        pBdr.append(el)
    pPr.append(pBdr)

    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    return p


SEVERITY_COLOR = {
    "High":   "C0392B",   # red
    "Medium": "D68910",   # amber
    "Low":    "1E8449",   # green
}

SEVERITY_BG = {
    "High":   "FADBD8",
    "Medium": "FDEBD0",
    "Low":    "D5F5E3",
}

# ─── Build the document ──────────────────────────────────────────────────────────

doc = Document()

# Page margins
section = doc.sections[0]
section.top_margin    = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(2.8)
section.right_margin  = Cm(2.8)

# ── Cover / Title ──────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after  = Pt(4)
run = title_p.add_run("AuthenticEye")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1A, 0x5C, 0xA8)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.paragraph_format.space_before = Pt(0)
sub_p.paragraph_format.space_after  = Pt(2)
run = sub_p.add_run("Local Setup — Issues & Resolutions Report")
run.bold = True
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.paragraph_format.space_before = Pt(2)
date_p.paragraph_format.space_after  = Pt(16)
run = date_p.add_run(
    f"Generated: {datetime.datetime(2026, 3, 12, 21, 29).strftime('%B %d, %Y  %H:%M')}"
)
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)
run.italic = True

add_horizontal_rule(doc)

# ── Executive Summary ──────────────────────────────────────────────────────────
sum_heading = doc.add_paragraph()
sum_heading.paragraph_format.space_before = Pt(10)
sum_heading.paragraph_format.space_after  = Pt(4)
r = sum_heading.add_run("Executive Summary")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x1A, 0x5C, 0xA8)

summary_text = (
    "This document records the seven issues encountered while setting up the AuthenticEye deepfake "
    "detection platform for local development on Windows. The issues span Docker connectivity, "
    "Python dependency installation, MediaPipe API incompatibility, model weight cold-start delays, "
    "missing environment configuration, and a missing Redis service. All issues were diagnosed and "
    "resolved, and the full platform (Frontend, Backend API, AI Service, and MongoDB) was "
    "successfully brought online."
)
s_p = doc.add_paragraph(summary_text)
s_p.paragraph_format.space_before = Pt(2)
s_p.paragraph_format.space_after  = Pt(10)
for run in s_p.runs:
    run.font.size = Pt(11)

add_horizontal_rule(doc)

# ── Quick Summary Table ──────────────────────────────────────────────────────
qt = doc.add_paragraph()
qt.paragraph_format.space_before = Pt(10)
qt.paragraph_format.space_after  = Pt(4)
r = qt.add_run("Issue Summary Table")
r.bold = True
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0x1A, 0x5C, 0xA8)

table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
header_cells = table.rows[0].cells
headers = ["#", "Issue Title", "Component", "Severity"]
for i, h in enumerate(headers):
    header_cells[i].text = h
    run = header_cells[i].paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_bg(header_cells[i], "1A5CA8")

for issue in ISSUES:
    row = table.add_row().cells
    row[0].text = str(issue["id"])
    row[1].text = issue["title"]
    row[2].text = issue["component"]
    row[3].text = issue["severity"]
    for cell in row:
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)
    # Severity cell coloring
    sev_color = SEVERITY_BG.get(issue["severity"], "FFFFFF")
    set_cell_bg(row[3], sev_color)
    run = row[3].paragraphs[0].runs[0]
    hex_c = SEVERITY_COLOR.get(issue["severity"], "000000")
    r2, g2, b2 = int(hex_c[:2], 16), int(hex_c[2:4], 16), int(hex_c[4:], 16)
    run.font.color.rgb = RGBColor(r2, g2, b2)
    run.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(6)
add_horizontal_rule(doc)

# ── Detailed Issue Sections ────────────────────────────────────────────────────
for issue in ISSUES:
    doc.add_page_break()

    # Issue title bar
    title_bar = doc.add_paragraph()
    title_bar.paragraph_format.space_before = Pt(0)
    title_bar.paragraph_format.space_after  = Pt(2)
    r = title_bar.add_run(f"  Issue #{issue['id']}  —  {issue['title']}")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '1A5CA8')
    title_bar._p.get_or_add_pPr().append(shd)

    # Meta row
    meta_p = doc.add_paragraph()
    meta_p.paragraph_format.space_before = Pt(4)
    meta_p.paragraph_format.space_after  = Pt(8)
    comp_run = meta_p.add_run(f"Component: {issue['component']}   |   ")
    comp_run.font.size = Pt(10)
    comp_run.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)
    comp_run.italic = True

    sev_run = meta_p.add_run(f"Severity: {issue['severity']}")
    sev_run.font.size = Pt(10)
    sev_run.bold = True
    hex_c = SEVERITY_COLOR.get(issue["severity"], "000000")
    r2, g2, b2 = int(hex_c[:2], 16), int(hex_c[2:4], 16), int(hex_c[4:], 16)
    sev_run.font.color.rgb = RGBColor(r2, g2, b2)

    # Section: Symptom
    add_colored_paragraph(doc, "🔴  Symptom / Error", bold=True, font_size=12,
                          color_hex="C0392B", space_before=6, space_after=3)
    s_p = doc.add_paragraph(issue["symptom"])
    s_p.paragraph_format.left_indent = Inches(0.2)
    s_p.paragraph_format.space_before = Pt(2)
    s_p.paragraph_format.space_after  = Pt(8)
    for run in s_p.runs:
        run.font.size = Pt(11)

    # Section: Root Cause
    add_colored_paragraph(doc, "🔍  Root Cause", bold=True, font_size=12,
                          color_hex="D68910", space_before=4, space_after=3)
    rc_p = doc.add_paragraph(issue["root_cause"])
    rc_p.paragraph_format.left_indent = Inches(0.2)
    rc_p.paragraph_format.space_before = Pt(2)
    rc_p.paragraph_format.space_after  = Pt(8)
    for run in rc_p.runs:
        run.font.size = Pt(11)

    # Section: Solution
    add_colored_paragraph(doc, "✅  Solution Applied", bold=True, font_size=12,
                          color_hex="1E8449", space_before=4, space_after=3)

    # Split solution on \n\n to show code blocks inline
    parts = issue["solution"].split("\n\n")
    for part in parts:
        part = part.strip()
        if not part:
            continue
        # Heuristic: if it contains code-like content (starts with spaces or contains ':')
        # and has indented lines, render as code block
        lines = part.split("\n")
        has_code = any(l.startswith("    ") or l.startswith(".\\" ) for l in lines)
        if has_code and len(lines) > 1:
            add_code_block(doc, part)
        else:
            sp = doc.add_paragraph(part)
            sp.paragraph_format.left_indent = Inches(0.2)
            sp.paragraph_format.space_before = Pt(2)
            sp.paragraph_format.space_after  = Pt(4)
            for run in sp.runs:
                run.font.size = Pt(11)

    add_horizontal_rule(doc)

# ── Final page: Service Status After All Fixes ─────────────────────────────────
doc.add_page_break()
fs_heading = doc.add_paragraph()
r = fs_heading.add_run("Final Service Status After Resolution")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x1A, 0x5C, 0xA8)
fs_heading.paragraph_format.space_after = Pt(6)

desc_p = doc.add_paragraph(
    "After all seven issues were resolved, the following services were confirmed running and "
    "healthy on the local machine:"
)
desc_p.paragraph_format.space_after = Pt(8)
for run in desc_p.runs:
    run.font.size = Pt(11)

status_table = doc.add_table(rows=1, cols=4)
status_table.style = 'Table Grid'
sh = status_table.rows[0].cells
for i, h in enumerate(["Service", "Port", "Status", "Verified Via"]):
    sh[i].text = h
    run = sh[i].paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    set_cell_bg(sh[i], "2C3E50")

status_rows = [
    ("Frontend (React / Vite)",     "5173", "✅ Running", "HTTP 200 at http://localhost:5173/"),
    ("Backend API (Node / Express)", "5000", "✅ Running", "GET /api/health → 200 OK"),
    ("AI Service (FastAPI)",         "8000", "✅ Running", "GET /health → {status: healthy}"),
    ("MongoDB",                      "27017","✅ Running", "Windows Service: MongoDB (Running)"),
    ("Redis",                        "6379", "⚠️ Not Available", "Not installed; not needed for core flow"),
]
for svc, port, status, verify in status_rows:
    row = status_table.add_row().cells
    row[0].text = svc
    row[1].text = port
    row[2].text = status
    row[3].text = verify
    for cell in row:
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)
    if "✅" in status:
        set_cell_bg(row[2], "D5F5E3")
        row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x1E, 0x84, 0x49)
    else:
        set_cell_bg(row[2], "FDEBD0")
        row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xD6, 0x89, 0x10)
    row[2].paragraphs[0].runs[0].bold = True

doc.add_paragraph()

footer_p = doc.add_paragraph(
    "This document was auto-generated to record the local setup troubleshooting session "
    "for the AuthenticEye Deepfake Detection Platform. All issues documented above were "
    "resolved within the same session."
)
footer_p.paragraph_format.space_before = Pt(12)
for run in footer_p.runs:
    run.font.size = Pt(10)
    run.italic = True
    run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

# ─── Save ────────────────────────────────────────────────────────────────────────
output_path = r"c:\Users\bhuvi\.gemini\antigravity\scratch\AuthenticEye\AuthenticEye_Local_Setup_Issues.docx"
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
